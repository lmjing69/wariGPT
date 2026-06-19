import os

from rag.pdf_loader import extract_pdf_text


class UnsupportedFileError(Exception):
    """Raised when a file type cannot (yet) be ingested as text."""


TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".log"}
DEFERRED_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp",  # images (vision: future)
    ".mp3", ".wav", ".m4a", ".ogg",                     # audio (future)
}


def _extract_docx(file_path):
    from docx import Document

    document = Document(file_path)

    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    # Include table cell text so tabular content is searchable too.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def file_kind(filename):
    """Return a coarse kind label used by the API/UI."""

    ext = os.path.splitext(filename or "")[1].lower()

    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in TEXT_EXTENSIONS:
        return "text"
    if ext in DEFERRED_EXTENSIONS:
        return "deferred"
    return "unknown"


def extract_text(file_path):
    """Extract plain text from a supported document.

    Raises UnsupportedFileError for image/audio (deferred) and unknown types.
    """

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_pdf_text(file_path)
    if ext == ".docx":
        return _extract_docx(file_path)
    if ext in TEXT_EXTENSIONS:
        return _extract_txt(file_path)
    if ext in DEFERRED_EXTENSIONS:
        raise UnsupportedFileError(
            f"{ext} files can be attached but aren't analyzed yet."
        )

    raise UnsupportedFileError(f"Unsupported file type: {ext or 'unknown'}")
